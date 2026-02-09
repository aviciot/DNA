"""
Template Files Management API

Endpoints for uploading and managing template reference documents (.docx files)
"""

import logging
import hashlib
import os
import json
from typing import List, Optional
from uuid import UUID, uuid4
from datetime import datetime
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from pydantic import BaseModel, Field

from ..database import get_db_pool
from ..auth import get_current_user, require_admin
from ..config import settings
from ..redis_client import redis_client

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/template-files", tags=["Template Files"])

# File storage configuration
UPLOAD_DIR = "/app/uploads/templates"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_MIME_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",  # Legacy .doc
]


# =============================================================================
# Pydantic Models
# =============================================================================

class TemplateFileBase(BaseModel):
    """Base model for template file"""
    original_filename: str
    description: Optional[str] = None
    version: Optional[str] = None
    notes: Optional[str] = None


class TemplateFileCreate(TemplateFileBase):
    """Model for creating a template file"""
    iso_standard_ids: List[UUID] = Field(default_factory=list, description="ISO standards this template applies to")


class TemplateFileResponse(BaseModel):
    """Response model for template file"""
    id: UUID
    filename: str
    original_filename: str
    file_path: str
    file_size_bytes: int
    file_hash: str
    mime_type: str
    description: Optional[str]
    version: Optional[str]
    notes: Optional[str]
    status: str
    uploaded_by: int
    uploaded_at: datetime
    updated_at: datetime

    # Computed fields
    built_templates_count: int = 0

    class Config:
        from_attributes = True


class TemplateFileListResponse(BaseModel):
    """Response model for listing template files"""
    id: UUID
    filename: str
    original_filename: str
    file_size_bytes: int
    description: Optional[str]
    status: str
    uploaded_at: datetime
    built_templates_count: int = 0


# =============================================================================
# Helper Functions
# =============================================================================

def compute_file_hash(file_content: bytes) -> str:
    """Compute SHA-256 hash of file content"""
    return hashlib.sha256(file_content).hexdigest()


def generate_unique_filename(original_filename: str, file_hash: str) -> str:
    """Generate unique filename using first 8 chars of hash + original name"""
    ext = os.path.splitext(original_filename)[1]
    base = os.path.splitext(original_filename)[0]
    # Sanitize filename
    safe_base = "".join(c for c in base if c.isalnum() or c in ('-', '_')).rstrip()
    return f"{file_hash[:8]}_{safe_base}{ext}"


def get_storage_path(filename: str, iso_codes: List[str]) -> tuple[str, str]:
    """
    Get storage path for file.

    Returns: (directory_path, full_file_path)
    """
    # If ISO codes provided, use first one for directory organization
    if iso_codes:
        # Use first ISO code, sanitize it (ISO 9001:2015 -> iso-9001-2015)
        iso_dir = iso_codes[0].lower().replace(" ", "-").replace(":", "-")
        directory = os.path.join(UPLOAD_DIR, iso_dir)
    else:
        directory = os.path.join(UPLOAD_DIR, "general")

    # Ensure directory exists
    os.makedirs(directory, exist_ok=True)

    full_path = os.path.join(directory, filename)
    return directory, full_path


# =============================================================================
# API Endpoints
# =============================================================================

@router.post("/upload", response_model=TemplateFileResponse, status_code=201)
async def upload_template_file(
    file: UploadFile = File(...),
    description: Optional[str] = Form(None),
    version: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    current_user: dict = Depends(require_admin)
):
    """
    Upload a template reference document (.docx).

    This endpoint stores the file and metadata WITHOUT triggering AI parsing.
    Use the /build endpoint later to parse and create templates from this file.

    NOTE: ISO standards are NOT assigned at upload time. They are assigned to
    individual templates in the Template Catalog after building.

    Requires: Admin role

    Form Data:
        - file: .docx file to upload (required)
        - description: Description of the reference document
        - version: Version identifier
        - notes: Admin notes

    Returns:
        Template file metadata
    """
    try:

        # Validate file type
        if file.content_type not in ALLOWED_MIME_TYPES:
            raise HTTPException(
                400,
                f"Invalid file type. Only Word documents (.docx) are supported. Got: {file.content_type}"
            )

        # Read file content
        file_content = await file.read()
        file_size = len(file_content)

        # Validate file size
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                400,
                f"File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024:.1f} MB"
            )

        if file_size == 0:
            raise HTTPException(400, "File is empty")

        # Compute file hash
        file_hash = compute_file_hash(file_content)

        pool = await get_db_pool()
        async with pool.acquire() as conn:
            # Check for duplicate by hash
            existing = await conn.fetchrow(
                f"SELECT id, original_filename FROM {settings.DATABASE_APP_SCHEMA}.template_files WHERE file_hash = $1 AND status = 'uploaded'",
                file_hash
            )

            if existing:
                raise HTTPException(
                    409,
                    f"This file already exists as '{existing['original_filename']}' (ID: {existing['id']})"
                )

            # Generate unique filename
            unique_filename = generate_unique_filename(file.filename, file_hash)

            # Get storage path (no ISO codes needed - store in general folder)
            directory = os.path.join(UPLOAD_DIR, "general")
            os.makedirs(directory, exist_ok=True)
            file_path = os.path.join(directory, unique_filename)

            # Save file to disk
            try:
                with open(file_path, 'wb') as f:
                    f.write(file_content)
                logger.info(f"Saved file to {file_path}")
            except Exception as e:
                logger.error(f"Failed to save file: {e}")
                raise HTTPException(500, f"Failed to save file: {str(e)}")

            # Insert into database
            try:
                row = await conn.fetchrow(f"""
                    INSERT INTO {settings.DATABASE_APP_SCHEMA}.template_files (
                        filename, original_filename, file_path, file_size_bytes,
                        file_hash, mime_type, description, version, notes,
                        status, uploaded_by
                    )
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, 'uploaded', $10)
                    RETURNING id, filename, original_filename, file_path, file_size_bytes,
                              file_hash, mime_type, description, version, notes,
                              status, uploaded_by, uploaded_at, updated_at
                """, unique_filename, file.filename, file_path, file_size,
                    file_hash, file.content_type, description, version, notes,
                    current_user.get('user_id'))

                file_id = row['id']

                logger.info(f"Uploaded template file {file_id} by user {current_user.get('user_id')}")

                # Build response
                result = dict(row)
                result['built_templates_count'] = 0

                return result

            except Exception as e:
                # Rollback: delete file from disk
                try:
                    os.remove(file_path)
                except:
                    pass
                logger.error(f"Database error during file upload: {e}")
                raise HTTPException(500, f"Failed to save file metadata: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error during file upload: {e}")
        raise HTTPException(500, f"File upload failed: {str(e)}")


@router.get("", response_model=List[TemplateFileListResponse])
async def list_template_files(
    status: Optional[str] = "uploaded",
    current_user: dict = Depends(get_current_user)
):
    """
    List all template reference files.

    Query Parameters:
        - status: Filter by status (default: 'uploaded'). Use 'all' for all statuses.

    Returns:
        List of template files with built template counts
    """
    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            # Use the view for efficient querying
            query = f"""
                SELECT
                    id, filename, original_filename, file_size_bytes,
                    description, status, uploaded_at,
                    built_templates_count
                FROM {settings.DATABASE_APP_SCHEMA}.v_template_files_with_details
                WHERE 1=1
            """

            params = []
            param_count = 1

            if status and status != "all":
                query += f" AND status = ${param_count}"
                params.append(status)
                param_count += 1

            query += " ORDER BY uploaded_at DESC"

            rows = await conn.fetch(query, *params)

            return [dict(row) for row in rows]

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing template files: {e}")
        raise HTTPException(500, f"Failed to list template files: {str(e)}")


@router.get("/{file_id}", response_model=TemplateFileResponse)
async def get_template_file(
    file_id: UUID,
    current_user: dict = Depends(get_current_user)
):
    """
    Get details of a specific template file.

    Returns:
        Template file details with statistics
    """
    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            # Get file details
            row = await conn.fetchrow(f"""
                SELECT
                    id, filename, original_filename, file_path, file_size_bytes,
                    file_hash, mime_type, description, version, notes,
                    status, uploaded_by, uploaded_at, updated_at
                FROM {settings.DATABASE_APP_SCHEMA}.template_files
                WHERE id = $1
            """, file_id)

            if not row:
                raise HTTPException(404, f"Template file {file_id} not found")

            # Get built templates count
            built_count = await conn.fetchval(f"""
                SELECT COUNT(*)
                FROM {settings.DATABASE_APP_SCHEMA}.catalog_templates
                WHERE template_file_id = $1
            """, file_id) or 0

            result = dict(row)
            result['built_templates_count'] = built_count

            return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting template file {file_id}: {e}")
        raise HTTPException(500, f"Failed to get template file: {str(e)}")


@router.delete("/{file_id}", status_code=204)
async def delete_template_file(
    file_id: UUID,
    current_user: dict = Depends(require_admin)
):
    """
    Delete a template file.

    Requires: Admin role

    Note: This will also delete all associated ISO mappings and built templates.
          Use with caution!
    """
    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            # Get file info
            row = await conn.fetchrow(
                f"SELECT file_path, status FROM {settings.DATABASE_APP_SCHEMA}.template_files WHERE id = $1",
                file_id
            )

            if not row:
                raise HTTPException(404, f"Template file {file_id} not found")

            file_path = row['file_path']

            # Delete from database (CASCADE will delete mappings)
            await conn.execute(
                f"DELETE FROM {settings.DATABASE_APP_SCHEMA}.template_files WHERE id = $1",
                file_id
            )

            # Delete file from disk
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Deleted file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to delete file from disk: {e}")

            logger.info(f"Deleted template file {file_id} by user {current_user.get('user_id')}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting template file {file_id}: {e}")
        raise HTTPException(500, f"Failed to delete template file: {str(e)}")


@router.get("/{file_id}/preview")
async def preview_template_file(
    file_id: UUID,
    current_user: dict = Depends(get_current_user)
):
    """
    Get HTML preview of a reference template file.

    Converts .docx to HTML for inline preview.

    Requires: Authenticated user
    """
    from fastapi.responses import HTMLResponse
    from docx import Document
    import html

    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            file_row = await conn.fetchrow(f"""
                SELECT id, original_filename, file_path
                FROM {settings.DATABASE_APP_SCHEMA}.template_files
                WHERE id = $1
            """, file_id)

            if not file_row:
                raise HTTPException(404, f"Template file {file_id} not found")

            file_path = file_row['file_path']

            if not os.path.exists(file_path):
                raise HTTPException(404, f"File not found on disk: {file_path}")

            # Convert .docx to HTML
            try:
                doc = Document(file_path)

                # Build HTML content
                html_parts = [
                    '<!DOCTYPE html>',
                    '<html>',
                    '<head>',
                    '<meta charset="UTF-8">',
                    '<style>',
                    'body { font-family: "Calibri", "Arial", sans-serif; margin: 0; padding: 0; background: #fff; color: #000; max-width: 100%; }',
                    'h1 { font-size: 24px; font-weight: bold; margin: 20px 0 10px 0; color: #1a1a1a; }',
                    'h2 { font-size: 20px; font-weight: bold; margin: 18px 0 8px 0; color: #1a1a1a; }',
                    'h3 { font-size: 16px; font-weight: bold; margin: 16px 0 6px 0; color: #1a1a1a; }',
                    'p { margin: 8px 0; line-height: 1.5; word-wrap: break-word; }',
                    'table { border-collapse: collapse; width: 100%; margin: 10px 0; table-layout: auto; }',
                    'td, th { border: 1px solid #ccc; padding: 8px; text-align: left; word-wrap: break-word; }',
                    'th { background-color: #f0f0f0; font-weight: bold; }',
                    '.document-title { font-size: 28px; font-weight: bold; text-align: center; margin: 30px 0; }',
                    '@media (max-width: 768px) { body { font-size: 14px; } h1 { font-size: 20px; } h2 { font-size: 18px; } }',
                    '</style>',
                    '</head>',
                    '<body>',
                ]

                # Add document title
                html_parts.append(f'<div class="document-title">{html.escape(file_row["original_filename"])}</div>')

                # Process paragraphs
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    # Determine heading level
                    style_name = para.style.name.lower()
                    if 'heading 1' in style_name or 'title' in style_name:
                        html_parts.append(f'<h1>{html.escape(text)}</h1>')
                    elif 'heading 2' in style_name:
                        html_parts.append(f'<h2>{html.escape(text)}</h2>')
                    elif 'heading 3' in style_name:
                        html_parts.append(f'<h3>{html.escape(text)}</h3>')
                    else:
                        html_parts.append(f'<p>{html.escape(text)}</p>')

                # Process tables
                for table in doc.tables:
                    html_parts.append('<table>')
                    for i, row in enumerate(table.rows):
                        html_parts.append('<tr>')
                        for cell in row.cells:
                            cell_text = html.escape(cell.text.strip())
                            if i == 0:  # First row as header
                                html_parts.append(f'<th>{cell_text}</th>')
                            else:
                                html_parts.append(f'<td>{cell_text}</td>')
                        html_parts.append('</tr>')
                    html_parts.append('</table>')

                html_parts.extend(['</body>', '</html>'])

                html_content = '\n'.join(html_parts)

                return HTMLResponse(content=html_content)

            except Exception as e:
                logger.error(f"Failed to convert document to HTML: {e}")
                raise HTTPException(500, f"Failed to preview document: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error previewing template file {file_id}: {e}")
        raise HTTPException(500, f"Failed to preview file: {str(e)}")


@router.get("/{file_id}/download")
async def download_template_file(
    file_id: UUID,
    current_user: dict = Depends(get_current_user)
):
    """
    Download a reference template file.

    Returns the .docx file for download.

    Requires: Authenticated user
    """
    from fastapi.responses import FileResponse

    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            file_row = await conn.fetchrow(f"""
                SELECT id, original_filename, file_path
                FROM {settings.DATABASE_APP_SCHEMA}.template_files
                WHERE id = $1
            """, file_id)

            if not file_row:
                raise HTTPException(404, f"Template file {file_id} not found")

            file_path = file_row['file_path']

            if not os.path.exists(file_path):
                raise HTTPException(404, f"File not found on disk: {file_path}")

            return FileResponse(
                path=file_path,
                filename=file_row['original_filename'],
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading template file {file_id}: {e}")
        raise HTTPException(500, f"Failed to download file: {str(e)}")


@router.post("/{file_id}/build", status_code=202)
async def build_template_from_file(
    file_id: UUID,
    current_user: dict = Depends(require_admin)
):
    """
    Build/parse a template from an uploaded file.

    Triggers async AI worker to parse the .docx file and create a template.

    Requires: Admin role

    Returns:
        202 Accepted with task_id for progress tracking
    """
    try:
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            # Get file details
            file_row = await conn.fetchrow(f"""
                SELECT id, filename, original_filename, file_path, status
                FROM {settings.DATABASE_APP_SCHEMA}.template_files
                WHERE id = $1
            """, file_id)

            if not file_row:
                raise HTTPException(404, f"Template file {file_id} not found")

            if file_row['status'] != 'uploaded':
                raise HTTPException(400, f"File status is '{file_row['status']}', must be 'uploaded'")

            # Check if file exists on disk
            file_path = file_row['file_path']
            if not os.path.exists(file_path):
                raise HTTPException(500, f"File not found on disk: {file_path}")

            # Check for existing in-progress task for this file
            existing_task = await conn.fetchrow(f"""
                SELECT id, status, created_at
                FROM {settings.DATABASE_APP_SCHEMA}.ai_tasks
                WHERE template_file_id = $1
                  AND status IN ('pending', 'processing')
                  AND created_at > NOW() - INTERVAL '10 minutes'
                LIMIT 1
            """, file_id)

            if existing_task:
                logger.info(f"File {file_id} already has in-progress task {existing_task['id']}")
                return {
                    "task_id": str(existing_task['id']),
                    "status": existing_task['status'],
                    "message": "Template build already in progress",
                    "created_at": existing_task['created_at'].isoformat()
                }

            # Get default parser LLM provider
            llm_row = await conn.fetchrow(f"""
                SELECT id, name, model
                FROM {settings.DATABASE_APP_SCHEMA}.llm_providers
                WHERE enabled = true AND is_default_parser = true
                LIMIT 1
            """)

            if not llm_row:
                raise HTTPException(500, "No default parser LLM provider configured")

            # Create AI task
            task_id = uuid4()
            trace_id = uuid4()

            await conn.execute(f"""
                INSERT INTO {settings.DATABASE_APP_SCHEMA}.ai_tasks (
                    id, task_type, status, progress, current_step,
                    llm_provider_id, llm_provider, llm_model,
                    template_file_id, trace_id, created_by
                )
                VALUES ($1, 'template_parse', 'pending', 0, 'Queued for processing',
                        $2, $3, $4, $5, $6, $7)
            """, task_id, llm_row['id'], llm_row['name'], llm_row['model'],
                file_id, trace_id, current_user.get('user_id'))

            # Queue task in Redis Stream
            task_message = {
                "task_id": str(task_id),
                "template_file_id": str(file_id),
                "file_path": file_path,
                "original_filename": file_row['original_filename'],
                "llm_provider": llm_row['name'],
                "created_by": str(current_user.get('user_id'))
            }

            # Add to Redis Stream using the correct method
            stream_id = await redis_client.add_to_stream(
                "template:parse",
                task_message
            )

            logger.info(f"Queued template build task {task_id} for file {file_id} (stream: {stream_id})")

            return {
                "task_id": str(task_id),
                "status": "pending",
                "message": "Template parsing queued. Use /ws/tasks/{task_id} to track progress.",
                "file_id": str(file_id),
                "filename": file_row['original_filename']
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error queueing build task for file {file_id}: {e}")
        raise HTTPException(500, f"Failed to queue build task: {str(e)}")
