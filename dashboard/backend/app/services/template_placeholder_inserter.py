"""
DNA Backend - Template Placeholder Inserter
============================================
Modifies .docx files to insert placeholders in place of original text.

When a template is created, this service:
1. Takes the original reference .docx file
2. Replaces identified text with {{placeholders}}
3. Saves as the template .docx file

This allows Fill & Preview to work by having actual placeholders in the document.
"""

import os
import logging
import re
from typing import List, Dict, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger(__name__)


class TemplatePlaceholderInserter:
    """Insert placeholders into .docx template files."""

    def __init__(self):
        """Initialize inserter."""
        pass

    def create_template_from_reference(
        self,
        reference_docx_path: str,
        replacements: List[Dict[str, Any]],
        output_path: str
    ) -> Dict[str, Any]:
        """
        Create template .docx by inserting placeholders into reference document.

        Args:
            reference_docx_path: Path to original reference .docx
            replacements: List of {original_text, placeholder} mappings
            output_path: Path to save template .docx

        Returns:
            Dict with:
            {
                "success": bool,
                "template_path": str,
                "replacements_made": int,
                "error": Optional[str]
            }
        """
        try:
            logger.info(f"Creating template from: {reference_docx_path}")
            logger.info(f"Replacements to make: {len(replacements)}")

            # Load reference document
            doc = Document(reference_docx_path)

            # Track replacements made
            replacements_made = 0

            # Process each replacement
            for replacement in replacements:
                original_text = replacement.get("original_text", "").strip()
                placeholder = replacement.get("placeholder", "").strip()

                if not original_text or not placeholder:
                    logger.warning(f"Skipping invalid replacement: {replacement}")
                    continue

                # Ensure placeholder has {{}} format
                if not placeholder.startswith("{{"):
                    placeholder = f"{{{{{placeholder}}}}}"

                logger.info(f"Replacing: '{original_text}' → '{placeholder}'")

                # Replace in all document parts
                count = 0
                count += self._replace_in_paragraphs(doc.paragraphs, original_text, placeholder)
                count += self._replace_in_tables(doc.tables, original_text, placeholder)
                count += self._replace_in_sections(doc, original_text, placeholder)

                if count > 0:
                    replacements_made += 1
                    logger.info(f"  Made {count} replacements for '{original_text}'")
                else:
                    logger.warning(f"  No occurrences found for '{original_text}'")

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save template
            doc.save(output_path)
            logger.info(f"Template saved: {output_path}")
            logger.info(f"Total replacements made: {replacements_made}/{len(replacements)}")

            return {
                "success": True,
                "template_path": output_path,
                "replacements_made": replacements_made,
                "total_expected": len(replacements),
                "error": None
            }

        except Exception as e:
            logger.error(f"Error creating template: {e}")
            return {
                "success": False,
                "template_path": None,
                "replacements_made": 0,
                "total_expected": len(replacements),
                "error": str(e)
            }

    def _replace_in_paragraphs(
        self,
        paragraphs: List[Paragraph],
        original_text: str,
        placeholder: str
    ) -> int:
        """Replace text in paragraphs. Returns count of replacements."""
        count = 0
        for para in paragraphs:
            if original_text in para.text:
                # Replace text while preserving formatting
                new_text = para.text.replace(original_text, placeholder)
                self._update_paragraph_text(para, new_text)
                count += para.text.count(placeholder)
        return count

    def _replace_in_tables(
        self,
        tables: List[Table],
        original_text: str,
        placeholder: str
    ) -> int:
        """Replace text in tables. Returns count of replacements."""
        count = 0
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    count += self._replace_in_paragraphs(
                        cell.paragraphs,
                        original_text,
                        placeholder
                    )
        return count

    def _replace_in_sections(
        self,
        doc: Document,
        original_text: str,
        placeholder: str
    ) -> int:
        """Replace text in headers and footers. Returns count of replacements."""
        count = 0
        for section in doc.sections:
            # Headers
            if section.header:
                count += self._replace_in_paragraphs(
                    section.header.paragraphs,
                    original_text,
                    placeholder
                )
                count += self._replace_in_tables(
                    section.header.tables,
                    original_text,
                    placeholder
                )

            # Footers
            if section.footer:
                count += self._replace_in_paragraphs(
                    section.footer.paragraphs,
                    original_text,
                    placeholder
                )
                count += self._replace_in_tables(
                    section.footer.tables,
                    original_text,
                    placeholder
                )
        return count

    def _update_paragraph_text(self, para: Paragraph, new_text: str):
        """Update paragraph text while preserving formatting."""
        # Clear existing runs
        for run in para.runs:
            run.text = ""

        # Add new text to first run (preserves some formatting)
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


def create_template_with_placeholders(
    reference_docx_path: str,
    replacements: List[Dict[str, Any]],
    output_path: str
) -> Dict[str, Any]:
    """
    Convenience function to create template with placeholders.

    Args:
        reference_docx_path: Path to original reference .docx
        replacements: List of {original_text, placeholder} mappings
        output_path: Path to save template .docx

    Returns:
        Dict with success status, path, and stats
    """
    inserter = TemplatePlaceholderInserter()
    return inserter.create_template_from_reference(
        reference_docx_path,
        replacements,
        output_path
    )


# Test
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python template_placeholder_inserter.py <reference.docx>")
        sys.exit(1)

    reference_path = sys.argv[1]
    test_replacements = [
        {"original_text": "ACME Corporation", "placeholder": "{{company_name}}"},
        {"original_text": "John Smith", "placeholder": "{{ciso_name}}"},
    ]

    output_path = "output/template_test.docx"

    result = create_template_with_placeholders(
        reference_path,
        test_replacements,
        output_path
    )

    if result["success"]:
        print(f"✅ Template created: {result['template_path']}")
        print(f"   Replacements: {result['replacements_made']}/{result['total_expected']}")
    else:
        print(f"❌ Error: {result['error']}")
