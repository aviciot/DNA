"""
DNA Backend - Template Parser Service
======================================
Uses LLM to parse Word documents and extract fillable fields.
"""

import anthropic
import json
import logging
from typing import Dict, List, Optional
from docx import Document
import os

logger = logging.getLogger(__name__)


class TemplateParser:
    """Parse document templates and identify fillable fields using Claude."""
    
    def __init__(self):
        """Initialize with Anthropic API key."""
        self.client = anthropic.Anthropic(
            api_key=os.getenv("ANTHROPIC_API_KEY", "")
        )
        self.model = "claude-sonnet-4-20250514"
    
    def extract_text_from_docx(self, file_path: str) -> Dict[str, any]:
        """
        Extract text and structure from Word document.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dict with paragraphs, tables, and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name,
                        "is_heading": para.style.name.startswith("Heading")
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from document: {e}")
            raise
    
    def parse_template_with_llm(self, document_content: Dict, document_name: str) -> Dict:
        """
        Use Claude to analyze document and identify fillable fields.
        
        Args:
            document_content: Extracted document structure
            document_name: Name of the document
            
        Returns:
            Dict with tagged template and field metadata
        """
        try:
            # Prepare document text for LLM
            full_text = "\n".join([p["text"] for p in document_content["paragraphs"]])
            
            prompt = f"""You are analyzing a certification document template called "{document_name}".

Your task is to:
1. Identify all fields that should be fillable (company names, dates, addresses, names, etc.)
2. Replace these fields with semantic tags in the format {{{{field_name}}}}
3. Create a metadata list of all fields with type and description

DOCUMENT CONTENT:
{full_text}

Please respond with JSON in this EXACT format:
{{
  "tagged_document": "The document text with {{{{tags}}}} replacing fillable fields",
  "fields": [
    {{
      "name": "company_name",
      "type": "text",
      "required": true,
      "description": "Legal name of the organization",
      "example": "Acme Corporation Ltd."
    }},
    {{
      "name": "audit_date",
      "type": "date",
      "required": true,
      "description": "Date of the certification audit",
      "example": "2026-03-15"
    }}
  ],
  "document_type": "policy|procedure|form|checklist"
}}

Field types can be: text, date, number, email, phone, address, textarea, select
Be thorough - identify ALL fillable fields including headers, footers, repeated information."""

            message = self.client.messages.create(
                model=self.model,
                max_tokens=8000,
                temperature=0,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Parse response
            response_text = message.content[0].text
            
            # Extract JSON from response (handle markdown code blocks)
            if "```json" in response_text:
                json_start = response_text.find("```json") + 7
                json_end = response_text.find("```", json_start)
                response_text = response_text[json_start:json_end].strip()
            elif "```" in response_text:
                json_start = response_text.find("```") + 3
                json_end = response_text.find("```", json_start)
                response_text = response_text[json_start:json_end].strip()
            
            result = json.loads(response_text)
            
            # Validate structure
            if "tagged_document" not in result or "fields" not in result:
                raise ValueError("LLM response missing required fields")
            
            logger.info(f"Successfully parsed template with {len(result['fields'])} fields")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing template with LLM: {e}")
            raise
    
    def parse_document_template(self, file_path: str, document_name: str) -> Dict:
        """
        Complete pipeline: extract text, parse with LLM.
        
        Args:
            file_path: Path to document file
            document_name: Name of the document
            
        Returns:
            Complete template structure with tags and metadata
        """
        logger.info(f"Parsing document template: {document_name}")
        
        # Step 1: Extract text from document
        document_content = self.extract_text_from_docx(file_path)
        
        # Step 2: Parse with LLM to identify fields
        parsed_template = self.parse_template_with_llm(document_content, document_name)
        
        # Step 3: Combine with original structure
        result = {
            "template_structure": {
                "tagged_content": parsed_template["tagged_document"],
                "original_paragraphs": document_content["paragraphs"],
                "tables": document_content["tables"]
            },
            "fields_metadata": parsed_template["fields"],
            "document_type": parsed_template.get("document_type", "document"),
            "parsing_metadata": {
                "total_fields": len(parsed_template["fields"]),
                "required_fields": sum(1 for f in parsed_template["fields"] if f.get("required", False)),
                "model_used": self.model
            }
        }
        
        logger.info(f"Template parsing complete: {result['parsing_metadata']}")
        return result


def calculate_document_completion(filled_data: Dict, fields_metadata: List[Dict]) -> float:
    """
    Calculate completion percentage based on filled fields.
    
    Args:
        filled_data: Dict of field_name -> filled_value
        fields_metadata: List of field definitions
        
    Returns:
        Completion percentage (0-100)
    """
    if not fields_metadata:
        return 0.0
    
    required_fields = [f for f in fields_metadata if f.get("required", False)]
    if not required_fields:
        required_fields = fields_metadata
    
    filled_count = sum(
        1 for field in required_fields 
        if filled_data.get(field["name"]) and str(filled_data[field["name"]]).strip()
    )
    
    return round((filled_count / len(required_fields)) * 100, 2)


def generate_filled_document(template_structure: Dict, filled_data: Dict) -> str:
    """
    Generate final document by replacing tags with filled data.
    
    Args:
        template_structure: Template with {{tags}}
        filled_data: Dict of tag -> value
        
    Returns:
        Final document text with values filled in
    """
    tagged_content = template_structure.get("tagged_content", "")
    
    # Replace all tags with filled data
    result = tagged_content
    for field_name, value in filled_data.items():
        tag = f"{{{{{field_name}}}}}"
        result = result.replace(tag, str(value))
    
    return result
