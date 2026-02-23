"""
DNA Backend - Placeholder Injector Service
==========================================
Injects placeholders into Word documents based on AI recommendations.

Takes a reference document and replacement map, creates a template with {{placeholders}}.
"""

import os
import logging
import re
from typing import Dict, List, Optional
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger(__name__)


class PlaceholderInjector:
    """Inject placeholders into Word documents."""

    def __init__(self):
        """Initialize injector."""
        self.doc: Optional[Document] = None
        self.replacement_map: Dict[str, str] = {}

    def inject_placeholders(
        self,
        source_path: str,
        output_path: str,
        replacement_map: Dict[str, str]
    ) -> Dict:
        """
        Inject placeholders into a Word document.

        Args:
            source_path: Path to source/reference DOCX
            replacement_map: Dict of original_text -> placeholder
                Example: {"CISO": "{{ciso_name}}", "the organization": "{{company_name}}"}
            output_path: Path to save template DOCX

        Returns:
            Dict with:
            {
                "success": bool,
                "template_path": str,
                "replacements_made": int,
                "error": Optional[str]
            }
        """
        try:
            logger.info(f"Loading document: {source_path}")

            # Load document
            self.doc = Document(source_path)
            self.replacement_map = replacement_map

            # Track replacements
            total_replacements = 0

            # Replace in all document parts
            total_replacements += self._replace_in_paragraphs(self.doc.paragraphs)
            total_replacements += self._replace_in_tables(self.doc.tables)
            total_replacements += self._replace_in_sections(self.doc)

            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            # Save template
            self.doc.save(output_path)
            logger.info(f"Template saved: {output_path} ({total_replacements} replacements)")

            return {
                "success": True,
                "template_path": output_path,
                "replacements_made": total_replacements,
                "error": None
            }

        except Exception as e:
            logger.error(f"Error injecting placeholders: {e}")
            return {
                "success": False,
                "template_path": None,
                "replacements_made": 0,
                "error": str(e)
            }

    def _replace_in_paragraphs(self, paragraphs) -> int:
        """Replace text in paragraphs."""
        replacements = 0
        for para in paragraphs:
            if not para.text.strip():
                continue

            replacements += self._process_paragraph(para)

        return replacements

    def _replace_in_tables(self, tables) -> int:
        """Replace text in tables."""
        replacements = 0
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replacements += self._replace_in_paragraphs(cell.paragraphs)
        return replacements

    def _replace_in_sections(self, doc: Document) -> int:
        """Replace text in headers and footers."""
        replacements = 0
        for section in doc.sections:
            # Headers
            if section.header:
                replacements += self._replace_in_paragraphs(section.header.paragraphs)
                replacements += self._replace_in_tables(section.header.tables)

            # Footers
            if section.footer:
                replacements += self._replace_in_paragraphs(section.footer.paragraphs)
                replacements += self._replace_in_tables(section.footer.tables)

        return replacements

    def _process_paragraph(self, para: Paragraph) -> int:
        """
        Process a paragraph and replace text with placeholders.

        Returns:
            Number of replacements made
        """
        original_text = para.text
        new_text = original_text
        replacements = 0

        # Apply each replacement from the map
        for original, placeholder in self.replacement_map.items():
            # Case-insensitive search, but preserve original case in non-matches
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            matches = pattern.findall(new_text)

            if matches:
                # Replace all occurrences
                new_text = pattern.sub(placeholder, new_text)
                replacements += len(matches)
                logger.debug(f"Replaced '{original}' with '{placeholder}' ({len(matches)} times)")

        # If any replacements were made, update the paragraph
        if new_text != original_text:
            self._replace_paragraph_text(para, new_text)

        return replacements

    def _replace_paragraph_text(self, para: Paragraph, new_text: str):
        """
        Replace paragraph text while preserving formatting.

        Note: This clears all runs and creates a single new run.
        For better formatting preservation, we could analyze runs,
        but for now we keep it simple.
        """
        # Clear existing runs
        for run in para.runs:
            run.text = ""

        # Add new text to first run (or create new run)
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


def inject_placeholders(
    source_path: str,
    output_path: str,
    replacement_map: Dict[str, str]
) -> Dict:
    """
    Convenience function to inject placeholders.

    Args:
        source_path: Path to source/reference DOCX
        output_path: Path to save template DOCX
        replacement_map: Dict of original_text -> placeholder

    Returns:
        Dict with success status and details
    """
    injector = PlaceholderInjector()
    return injector.inject_placeholders(source_path, output_path, replacement_map)


# Test/Demo
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python placeholder_injector.py <reference.docx>")
        sys.exit(1)

    source_path = sys.argv[1]

    # Test replacement map
    test_map = {
        "CISO": "{{ciso_name}}",
        "the organization": "{{company_name}}",
        "Company": "{{company_name}}",
    }

    print(f"\n{'='*60}")
    print(f"Injecting Placeholders into: {source_path}")
    print(f"{'='*60}\n")

    result = inject_placeholders(
        source_path,
        "output/template_with_placeholders.docx",
        test_map
    )

    if result["success"]:
        print(f"✅ Success! Made {result['replacements_made']} replacements")
        print(f"Template: {result['template_path']}")
    else:
        print(f"❌ Error: {result['error']}")
