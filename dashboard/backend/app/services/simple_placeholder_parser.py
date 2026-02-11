"""
DNA Backend - Simple Placeholder Parser
========================================
Extracts {{placeholder}} patterns from DOCX templates.

Part of Phase 1: Template Preview System (Proof of Concept)
"""

import re
import logging
from typing import Dict, List, Optional, Set, Tuple
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class PlaceholderParser:
    """Parse DOCX templates to extract {{placeholder}} fields."""

    # Regex pattern to match {{field_name}} or {{field_name|type|required}}
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')

    def __init__(self):
        """Initialize parser."""
        self.placeholders: Dict[str, Dict] = {}

    def parse_template(self, docx_path: str) -> Dict:
        """
        Main entry point: Parse template and extract all placeholders.

        Args:
            docx_path: Path to DOCX template file

        Returns:
            Dict with:
            {
                "template_path": str,
                "filename": str,
                "total_placeholders": int,
                "fields": [
                    {
                        "name": str,
                        "label": str,
                        "type": str,
                        "required": bool,
                        "evidence": bool,
                        "evidence_filename": Optional[str],
                        "locations": List[str]
                    }
                ]
            }
        """
        try:
            logger.info(f"Parsing template: {docx_path}")

            # Load document
            doc = Document(docx_path)

            # Reset placeholders
            self.placeholders = {}

            # Extract from different document parts
            self._extract_from_paragraphs(doc.paragraphs, "body")
            self._extract_from_tables(doc.tables)
            self._extract_from_sections(doc)

            # Convert to structured output
            fields = self._build_field_list()

            result = {
                "template_path": docx_path,
                "filename": docx_path.split("/")[-1].split("\\")[-1],
                "total_placeholders": len(fields),
                "fields": fields
            }

            logger.info(f"Found {len(fields)} placeholders")
            return result

        except Exception as e:
            logger.error(f"Error parsing template: {e}")
            raise

    def _extract_from_paragraphs(
        self,
        paragraphs: List[Paragraph],
        location_prefix: str
    ):
        """Extract placeholders from paragraphs."""
        for idx, para in enumerate(paragraphs):
            if not para.text.strip():
                continue

            matches = self.PLACEHOLDER_PATTERN.findall(para.text)
            for match in matches:
                location = f"{location_prefix}:paragraph:{idx}"
                self._process_placeholder(match, location)

    def _extract_from_tables(self, tables: List[Table]):
        """Extract placeholders from tables."""
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue

                        matches = self.PLACEHOLDER_PATTERN.findall(para.text)
                        for match in matches:
                            location = f"table:{table_idx}:row:{row_idx}:col:{col_idx}"
                            self._process_placeholder(match, location)

    def _extract_from_sections(self, doc: Document):
        """Extract placeholders from headers and footers."""
        for section_idx, section in enumerate(doc.sections):
            # Headers
            if section.header:
                self._extract_from_paragraphs(
                    section.header.paragraphs,
                    f"header:{section_idx}"
                )
                self._extract_from_tables(section.header.tables)

            # Footers
            if section.footer:
                self._extract_from_paragraphs(
                    section.footer.paragraphs,
                    f"footer:{section_idx}"
                )
                self._extract_from_tables(section.footer.tables)

    def _process_placeholder(self, match: str, location: str):
        """
        Process a single placeholder match.

        Formats supported:
        - {{field_name}}
        - {{field_name|type}}
        - {{field_name|type|required}}
        - {{field_name|type|required|evidence:filename}}
        """
        parts = match.split("|")
        field_name = parts[0].strip()

        # Parse metadata from explicit markup
        field_type = "text"
        required = False
        evidence = False
        evidence_filename = None

        if len(parts) > 1:
            field_type = parts[1].strip()

        if len(parts) > 2 and parts[2].strip().lower() == "required":
            required = True

        if len(parts) > 3 and parts[3].strip().startswith("evidence:"):
            evidence = True
            evidence_filename = parts[3].strip().split(":", 1)[1]

        # If field type not explicitly set, try to guess from name
        if field_type == "text":
            field_type = self._guess_field_type(field_name)

        # Add or update placeholder info
        if field_name not in self.placeholders:
            self.placeholders[field_name] = {
                "name": field_name,
                "label": self._generate_label(field_name),
                "type": field_type,
                "required": required,
                "evidence": evidence,
                "evidence_filename": evidence_filename,
                "locations": []
            }

        # Update required flag (if ANY location marks it required)
        if required:
            self.placeholders[field_name]["required"] = True

        # Add location
        self.placeholders[field_name]["locations"].append(location)

    def _guess_field_type(self, field_name: str) -> str:
        """
        Guess field type from field name.

        Examples:
        - "effective_date" → "date"
        - "ciso_email" → "email"
        - "logo" → "image"
        - "backup_solution" → "textarea"
        """
        name_lower = field_name.lower()

        # Date fields
        if any(keyword in name_lower for keyword in ["date", "time", "deadline", "due"]):
            return "date"

        # Email fields
        if "email" in name_lower or "mail" in name_lower:
            return "email"

        # Phone fields
        if "phone" in name_lower or "tel" in name_lower or "mobile" in name_lower:
            return "phone"

        # Image fields
        if any(keyword in name_lower for keyword in ["logo", "image", "photo", "picture"]):
            return "image"

        # File/evidence fields
        if any(keyword in name_lower for keyword in ["file", "document", "attachment", "evidence"]):
            return "file"

        # Number fields
        if any(keyword in name_lower for keyword in ["count", "number", "qty", "quantity", "amount"]):
            return "number"

        # Textarea fields (longer content)
        if any(keyword in name_lower for keyword in [
            "description", "details", "notes", "comment", "summary",
            "explanation", "justification", "rationale", "solution"
        ]):
            return "textarea"

        # Select fields (common choices)
        if any(keyword in name_lower for keyword in ["status", "priority", "level", "type", "category"]):
            return "select"

        # Default to text
        return "text"

    def _generate_label(self, field_name: str) -> str:
        """
        Generate human-readable label from field name.

        Examples:
        - "company_name" → "Company Name"
        - "ciso_email" → "CISO Email"
        - "backup_solution" → "Backup Solution"
        """
        # Split on underscores
        words = field_name.split("_")

        # Capitalize each word
        label = " ".join(word.capitalize() for word in words)

        # Handle common acronyms
        acronyms = ["ISO", "CISO", "IT", "HR", "CEO", "CTO", "CFO", "AWS", "API", "URL"]
        for acronym in acronyms:
            if acronym.lower() in [w.lower() for w in words]:
                label = label.replace(acronym.capitalize(), acronym)

        return label

    def _build_field_list(self) -> List[Dict]:
        """Convert internal placeholder dict to sorted field list."""
        fields = list(self.placeholders.values())

        # Sort by priority:
        # 1. Required fields first
        # 2. Then by usage count (more locations = higher priority)
        # 3. Then by type (images and dates first)
        # 4. Then alphabetically

        type_priority = {
            "image": 0,
            "date": 1,
            "email": 2,
            "text": 3,
            "textarea": 4,
            "select": 5,
            "number": 6,
            "phone": 7,
            "file": 8
        }

        fields.sort(key=lambda f: (
            not f["required"],           # Required first
            -len(f["locations"]),        # More usage = higher priority
            type_priority.get(f["type"], 99),  # Type priority
            f["name"]                    # Alphabetical
        ))

        return fields


def extract_placeholders_from_docx(docx_path: str) -> Dict:
    """
    Convenience function to parse template.

    Args:
        docx_path: Path to DOCX template

    Returns:
        Dict with template info and field list
    """
    parser = PlaceholderParser()
    return parser.parse_template(docx_path)


# Test/Demo
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python simple_placeholder_parser.py <template.docx>")
        sys.exit(1)

    template_path = sys.argv[1]

    print(f"\n{'='*60}")
    print(f"Parsing Template: {template_path}")
    print(f"{'='*60}\n")

    result = extract_placeholders_from_docx(template_path)

    print(f"Total Placeholders Found: {result['total_placeholders']}\n")

    for field in result['fields']:
        print(f"Field: {field['name']}")
        print(f"  Label: {field['label']}")
        print(f"  Type: {field['type']}")
        print(f"  Required: {field['required']}")
        if field['evidence']:
            print(f"  Evidence: {field['evidence_filename']}")
        print(f"  Locations: {len(field['locations'])} places")
        print()
