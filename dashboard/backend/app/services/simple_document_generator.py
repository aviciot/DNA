"""
DNA Backend - Simple Document Generator
========================================
Generates filled documents from templates with placeholder replacement.

Part of Phase 1: Template Preview System (Proof of Concept)
"""

import os
import re
import logging
import shutil
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate filled documents from templates."""

    def __init__(self):
        """Initialize generator."""
        self.template_doc: Optional[Document] = None
        self.filled_data: Dict[str, Any] = {}

    def generate_document(
        self,
        template_path: str,
        filled_data: Dict[str, Any],
        output_dir: str,
        output_filename: Optional[str] = None
    ) -> Dict[str, str]:
        """
        Generate filled document from template.

        Args:
            template_path: Path to template DOCX
            filled_data: Dict of placeholder -> value
            output_dir: Directory to save outputs
            output_filename: Optional custom filename (without extension)

        Returns:
            Dict with:
            {
                "success": bool,
                "docx_path": str,
                "pdf_path": str,
                "error": Optional[str]
            }
        """
        try:
            logger.info(f"Generating document from template: {template_path}")

            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)

            # Load template
            self.template_doc = Document(template_path)
            self.filled_data = filled_data

            # Replace placeholders in all document parts
            self._replace_in_paragraphs(self.template_doc.paragraphs)
            self._replace_in_tables(self.template_doc.tables)
            self._replace_in_sections(self.template_doc)

            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"filled_document_{timestamp}"

            # Save filled DOCX
            docx_path = os.path.join(output_dir, f"{output_filename}.docx")
            self.template_doc.save(docx_path)
            logger.info(f"Saved filled DOCX: {docx_path}")

            # Convert to PDF
            pdf_path = os.path.join(output_dir, f"{output_filename}.pdf")
            self._convert_to_pdf(docx_path, pdf_path)
            logger.info(f"Converted to PDF: {pdf_path}")

            return {
                "success": True,
                "docx_path": docx_path,
                "pdf_path": pdf_path,
                "error": None
            }

        except Exception as e:
            logger.error(f"Error generating document: {e}")
            return {
                "success": False,
                "docx_path": None,
                "pdf_path": None,
                "error": str(e)
            }

    def _replace_in_paragraphs(self, paragraphs):
        """Replace placeholders in paragraphs."""
        for para in paragraphs:
            if not para.text.strip():
                continue

            # Check if paragraph contains any placeholders
            if "{{" in para.text and "}}" in para.text:
                self._process_paragraph(para)

    def _replace_in_tables(self, tables):
        """Replace placeholders in tables."""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_paragraphs(cell.paragraphs)

    def _replace_in_sections(self, doc: Document):
        """Replace placeholders in headers and footers."""
        for section in doc.sections:
            # Headers
            if section.header:
                self._replace_in_paragraphs(section.header.paragraphs)
                self._replace_in_tables(section.header.tables)

            # Footers
            if section.footer:
                self._replace_in_paragraphs(section.footer.paragraphs)
                self._replace_in_tables(section.footer.tables)

    def _process_paragraph(self, para: Paragraph):
        """
        Process a paragraph containing placeholders.

        Handles both text and image placeholders.
        """
        original_text = para.text

        # Find all placeholders in this paragraph
        placeholder_pattern = re.compile(r'\{\{([^}]+)\}\}')
        matches = placeholder_pattern.findall(original_text)

        if not matches:
            return

        # Check if this is an image placeholder
        for match in matches:
            field_name = match.split("|")[0].strip()

            # Determine if this is an image field
            is_image = (
                "|image|" in match.lower() or
                "logo" in field_name.lower() or
                "image" in field_name.lower() or
                "photo" in field_name.lower()
            )

            if is_image and field_name in self.filled_data:
                # Replace paragraph with image
                image_path = self.filled_data[field_name]
                if image_path and os.path.exists(image_path):
                    self._insert_image(para, image_path)
                    return  # Image inserted, done with this paragraph

        # If not image, replace text placeholders
        new_text = original_text
        for match in matches:
            field_name = match.split("|")[0].strip()
            placeholder = f"{{{{{match}}}}}"

            if field_name in self.filled_data:
                value = self._format_value(field_name, self.filled_data[field_name], match)
                new_text = new_text.replace(placeholder, value)
            else:
                # Leave placeholder if no data provided
                logger.warning(f"No data provided for placeholder: {field_name}")

        # Replace text while preserving formatting
        if new_text != original_text:
            self._replace_paragraph_text(para, new_text)

    def _replace_paragraph_text(self, para: Paragraph, new_text: str):
        """Replace paragraph text while preserving formatting."""
        # Clear existing runs
        for run in para.runs:
            run.text = ""

        # Add new text to first run (preserves some formatting)
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)

    def _insert_image(self, para: Paragraph, image_path: str):
        """Insert image into paragraph, replacing text."""
        # Clear paragraph text
        for run in para.runs:
            run.text = ""

        # Add image
        try:
            # Determine image size (max width 2 inches)
            run = para.runs[0] if para.runs else para.add_run()
            run.add_picture(image_path, width=Inches(2.0))
            logger.info(f"Inserted image: {image_path}")
        except Exception as e:
            logger.error(f"Error inserting image {image_path}: {e}")
            # Fallback: add image path as text
            para.add_run(f"[Image: {os.path.basename(image_path)}]")

    def _format_value(self, field_name: str, value: Any, match: str) -> str:
        """
        Format value based on field type.

        Args:
            field_name: Name of the field
            value: Raw value
            match: Full match string (may contain type info)

        Returns:
            Formatted string value
        """
        # Check if type is specified in placeholder
        parts = match.split("|")
        field_type = parts[1].strip() if len(parts) > 1 else "text"

        # Format based on type
        if field_type == "date" or "date" in field_name.lower():
            return self._format_date(value)
        elif field_type == "number" or isinstance(value, (int, float)):
            return str(value)
        elif field_type == "email":
            return str(value).lower()
        else:
            return str(value)

    def _format_date(self, value: Any) -> str:
        """Format date value."""
        if isinstance(value, datetime):
            return value.strftime("%B %d, %Y")  # e.g., "February 11, 2026"
        elif isinstance(value, str):
            # Try to parse string date
            try:
                # Common formats
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%Y-%m-%d %H:%M:%S"]:
                    try:
                        dt = datetime.strptime(value, fmt)
                        return dt.strftime("%B %d, %Y")
                    except ValueError:
                        continue
                # If no format matches, return as-is
                return value
            except Exception:
                return value
        else:
            return str(value)

    def _convert_to_pdf(self, docx_path: str, pdf_path: str):
        """
        Convert DOCX to PDF.

        Uses different methods based on platform:
        - Windows: docx2pdf
        - Linux/Mac: LibreOffice or pypandoc
        """
        try:
            # Try docx2pdf first (Windows)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                logger.info("Converted to PDF using docx2pdf")
                return
            except ImportError:
                logger.debug("docx2pdf not available")
            except Exception as e:
                logger.warning(f"docx2pdf conversion failed: {e}")

            # Try LibreOffice (cross-platform)
            try:
                import subprocess

                # Find LibreOffice executable
                libreoffice_paths = [
                    "libreoffice",  # Linux
                    "/usr/bin/libreoffice",  # Linux
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # Mac
                    "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
                ]

                libreoffice_cmd = None
                for path in libreoffice_paths:
                    if shutil.which(path) or os.path.exists(path):
                        libreoffice_cmd = path
                        break

                if libreoffice_cmd:
                    output_dir = os.path.dirname(pdf_path)
                    subprocess.run([
                        libreoffice_cmd,
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", output_dir,
                        docx_path
                    ], check=True, capture_output=True)
                    logger.info("Converted to PDF using LibreOffice")
                    return
                else:
                    logger.debug("LibreOffice not found")
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}")

            # Try pypandoc (fallback)
            try:
                import pypandoc
                pypandoc.convert_file(
                    docx_path,
                    'pdf',
                    outputfile=pdf_path,
                    extra_args=['--pdf-engine=xelatex']
                )
                logger.info("Converted to PDF using pypandoc")
                return
            except ImportError:
                logger.debug("pypandoc not available")
            except Exception as e:
                logger.warning(f"pypandoc conversion failed: {e}")

            # If all methods fail, create a placeholder PDF message
            logger.error("No PDF conversion method available")
            raise Exception(
                "PDF conversion failed. Please install one of: "
                "docx2pdf (Windows), LibreOffice (any OS), or pypandoc"
            )

        except Exception as e:
            logger.error(f"Error converting to PDF: {e}")
            raise


def generate_filled_document(
    template_path: str,
    filled_data: Dict[str, Any],
    output_dir: str,
    output_filename: Optional[str] = None
) -> Dict[str, str]:
    """
    Convenience function to generate filled document.

    Args:
        template_path: Path to template DOCX
        filled_data: Dict of placeholder -> value
        output_dir: Directory to save outputs
        output_filename: Optional custom filename (without extension)

    Returns:
        Dict with success status, paths, and any errors
    """
    generator = DocumentGenerator()
    return generator.generate_document(
        template_path,
        filled_data,
        output_dir,
        output_filename
    )


# Test/Demo
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python simple_document_generator.py <template.docx>")
        sys.exit(1)

    template_path = sys.argv[1]

    # Test data
    test_data = {
        "company_name": "Acme Corporation",
        "effective_date": "2026-02-11",
        "ciso_name": "John Smith",
        "backup_solution": "We use AWS S3 for daily automated backups with 30-day retention and versioning enabled.",
        # "logo": "path/to/logo.png"  # Uncomment if you have a test logo
    }

    print(f"\n{'='*60}")
    print(f"Generating Document from Template: {template_path}")
    print(f"{'='*60}\n")

    result = generate_filled_document(
        template_path,
        test_data,
        output_dir="output",
        output_filename="test_document"
    )

    if result["success"]:
        print("✅ Document generated successfully!")
        print(f"DOCX: {result['docx_path']}")
        print(f"PDF: {result['pdf_path']}")
    else:
        print(f"❌ Error: {result['error']}")
