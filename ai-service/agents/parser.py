"""
Template Parser Agent (DEPRECATED)
===================================

⚠️ DEPRECATED: This file is kept for reference only.
Use agents/template.py (TemplateAgent) instead.

The new TemplateAgent extends BaseAgent and provides:
- Rate-limited LLM calls (prevents API overuse)
- Integrated telemetry
- Stateless design (thread-safe)
- Better error handling
- Template editing capability

Old code below (DO NOT USE):
----------------------------

Parses Word documents into structured ISO certification templates using Claude AI.

User Experience Focus:
- Clear progress updates at each step
- Detailed extraction of sections, fields, and requirements
- Validation against ISO standards
- Helpful error messages
"""

import logging
import time
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from anthropic import AsyncAnthropic
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

# Import telemetry for LLM call tracking
try:
    from telemetry import telemetry
    TELEMETRY_AVAILABLE = True
except ImportError:
    TELEMETRY_AVAILABLE = False
    logger.warning("Telemetry not available - skipping telemetry logging")


class TemplateParserAgent:
    """
    AI-powered parser for ISO certification Word documents.

    Extracts:
    - Document structure (sections, subsections)
    - Fillable fields with types and constraints
    - ISO standard requirements
    - Field relationships and dependencies
    """

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-5-20250929"):
        """
        Initialize parser agent.

        Args:
            api_key: Anthropic API key
            model: Claude model to use
        """
        self.client = AsyncAnthropic(api_key=api_key)
        self.model = model
        self.max_tokens = 4096

    def _extract_json_from_response(self, text: str) -> str:
        """
        Extract JSON from Claude's response, handling markdown code fences and extra text.

        Args:
            text: Raw response text from Claude

        Returns:
            Extracted JSON string
        """
        # Remove markdown code fences if present
        if "```json" in text:
            # Extract content between ```json and ```
            start = text.find("```json") + 7
            end = text.find("```", start)
            if end != -1:
                text = text[start:end].strip()
        elif "```" in text:
            # Extract content between ``` and ```
            start = text.find("```") + 3
            end = text.find("```", start)
            if end != -1:
                text = text[start:end].strip()

        # Try to find JSON object or array
        text = text.strip()

        # Find first { or [ and last } or ]
        start_brace = text.find("{")
        start_bracket = text.find("[")

        if start_brace == -1 and start_bracket == -1:
            raise ValueError("No JSON object or array found in response")

        if start_brace == -1:
            start = start_bracket
            end_char = "]"
        elif start_bracket == -1:
            start = start_brace
            end_char = "}"
        else:
            start = min(start_brace, start_bracket)
            end_char = "}" if start == start_brace else "]"

        # Find matching closing bracket/brace
        end = text.rfind(end_char)

        if end == -1 or end < start:
            raise ValueError(f"No matching {end_char} found for JSON")

        json_str = text[start:end + 1]
        return json_str

    async def parse_document(
        self,
        file_path: str,
        custom_rules: Optional[str] = None,
        iso_standard: Optional[str] = None,
        trace_id: Optional[str] = None,
        task_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse Word document into structured template.

        Args:
            file_path: Path to Word document
            custom_rules: Optional custom parsing rules
            iso_standard: Optional ISO standard (e.g., "ISO 9001:2015")
            trace_id: Optional trace ID for telemetry
            task_id: Optional task ID for telemetry

        Returns:
            Parsed template structure with sections, fields, and metadata
        """
        logger.info(f"Starting document parsing: {file_path}")

        # Store for telemetry in LLM calls
        self._current_trace_id = trace_id
        self._current_task_id = task_id

        # Telemetry: Agent started
        if TELEMETRY_AVAILABLE and trace_id and task_id:
            telemetry.agent_started(
                agent_name="TemplateAgent",
                trace_id=trace_id,
                task_id=task_id,
                file_path=file_path,
                iso_standard=iso_standard,
                has_custom_rules=bool(custom_rules)
            )

        agent_start_time = time.time()

        # Step 1: Load and extract document content
        logger.info("Step 1/4: Loading Word document...")
        doc_content = await self._extract_document_content(file_path)

        # Step 2: Analyze structure with Claude
        logger.info("Step 2/4: Analyzing document structure with Claude...")
        structure = await self._analyze_structure(doc_content, iso_standard, custom_rules)

        # Step 3: Extract fillable fields
        logger.info("Step 3/4: Extracting fillable fields...")
        fields = await self._extract_fields(doc_content, structure)

        # Step 4: Validate and enrich
        logger.info("Step 4/4: Validating and enriching template...")
        template = await self._validate_and_enrich(structure, fields, iso_standard)

        logger.info(f"Parsing complete: {len(template['sections'])} sections, {len(template['fields'])} fields")

        # Telemetry: Agent completed
        agent_duration = int(time.time() - agent_start_time)
        if TELEMETRY_AVAILABLE and trace_id and task_id:
            telemetry.agent_completed(
                agent_name="TemplateAgent",
                trace_id=trace_id,
                task_id=task_id,
                duration_seconds=agent_duration,
                result_summary={
                    "sections": len(template['sections']),
                    "fields": len(template['fields']),
                    "required_fields": template['metadata']['required_fields'],
                    "completion_estimate_minutes": template['metadata']['completion_estimate_minutes']
                }
            )

        return template

    async def _extract_document_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from Word document.

        Returns:
            Dictionary with paragraphs, tables, and metadata
        """
        try:
            # Check file exists
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")

            # Load document
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal",
                        "level": self._get_heading_level(para)
                    })

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "Untitled",
                "author": core_props.author or "Unknown",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }

            logger.info(f"Extracted: {len(paragraphs)} paragraphs, {len(tables)} tables")

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"Failed to extract document content: {e}")
            raise

    def _get_heading_level(self, paragraph) -> int:
        """Get heading level from paragraph style."""
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "heading 1" in style_name:
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        elif "heading" in style_name:
            # Extract number from "Heading 4", "Heading 5", etc.
            try:
                return int(style_name.split()[-1])
            except:
                return 0
        return 0

    async def _analyze_structure(
        self,
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str]
    ) -> Dict[str, Any]:
        """
        Use Claude to analyze document structure.

        Returns:
            Hierarchical structure with sections and subsections
        """
        # Build prompt
        prompt = self._build_structure_prompt(doc_content, iso_standard, custom_rules)

        # Call Claude
        try:
            # Telemetry: LLM request
            llm_start_time = time.time()
            if TELEMETRY_AVAILABLE and self._current_trace_id and self._current_task_id:
                telemetry.llm_request(
                    provider="anthropic",
                    model=self.model,
                    trace_id=self._current_trace_id,
                    task_id=self._current_task_id,
                    prompt_type="structure_analysis"
                )

            response = await self.client.messages.create(
                model=self.model,
                max_tokens=self.max_tokens,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # Telemetry: LLM response
            llm_duration_ms = int((time.time() - llm_start_time) * 1000)
            input_tokens = response.usage.input_tokens
            output_tokens = response.usage.output_tokens
            # Claude Sonnet 4.5 pricing: $3/M input, $15/M output
            cost_usd = (input_tokens / 1_000_000 * 3.0) + (output_tokens / 1_000_000 * 15.0)

            if TELEMETRY_AVAILABLE and self._current_trace_id and self._current_task_id:
                telemetry.llm_response(
                    provider="anthropic",
                    model=self.model,
                    trace_id=self._current_trace_id,
                    task_id=self._current_task_id,
                    duration_ms=llm_duration_ms,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    cost_usd=cost_usd
                )

            # Extract JSON from response
            content = response.content[0].text

            # Extract JSON (handles markdown code fences)
            import json
            json_str = self._extract_json_from_response(content)
            structure = json.loads(json_str)

            logger.info(f"Structure analyzed: {len(structure.get('sections', []))} top-level sections")

            return structure

        except Exception as e:
            logger.error(f"Failed to analyze structure: {e}")
            raise

    def _build_structure_prompt(
        self,
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str]
    ) -> str:
        """Build prompt for structure analysis."""

        # Format paragraphs for context
        paragraphs_text = "\n\n".join([
            f"[Level {p['level']}] {p['text']}" for p in doc_content['paragraphs'][:100]  # Limit to first 100
        ])

        prompt = f"""You are an expert at analyzing ISO certification templates and Word documents.

Analyze the following document and extract its hierarchical structure.

Document Metadata:
- Title: {doc_content['metadata']['title']}
- Paragraphs: {doc_content['metadata']['paragraph_count']}
- Tables: {doc_content['metadata']['table_count']}

ISO Standard: {iso_standard or 'Not specified'}

Document Content (excerpt):
{paragraphs_text}

Your task:
1. Identify all major sections and subsections
2. Determine the hierarchy (parent-child relationships)
3. Extract section titles and descriptions
4. Identify which sections contain fillable fields (forms, questionnaires, checklists)

{f'Custom Rules: {custom_rules}' if custom_rules else ''}

Return ONLY valid JSON in this exact format:
{{
  "sections": [
    {{
      "id": "section_1",
      "title": "Company Information",
      "level": 1,
      "order": 1,
      "description": "Basic company details and contact information",
      "has_fields": true,
      "subsections": [
        {{
          "id": "section_1_1",
          "title": "General Details",
          "level": 2,
          "order": 1,
          "description": "Company name, address, registration",
          "has_fields": true
        }}
      ]
    }}
  ]
}}

Important:
- Use descriptive IDs (e.g., "section_company_info")
- Include ALL sections found in the document
- Set has_fields=true for sections containing forms/questionnaires
- Maintain correct parent-child hierarchy
"""

        return prompt

    async def _extract_fields(
        self,
        doc_content: Dict[str, Any],
        structure: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Extract fillable fields from document.

        Returns:
            List of field definitions with types and constraints
        """
        # Build prompt for field extraction
        prompt = self._build_fields_prompt(doc_content, structure)

        try:
            # Telemetry: LLM request
            llm_start_time = time.time()
            if TELEMETRY_AVAILABLE and self._current_trace_id and self._current_task_id:
                telemetry.llm_request(
                    provider="anthropic",
                    model=self.model,
                    trace_id=self._current_trace_id,
                    task_id=self._current_task_id,
                    prompt_type="field_extraction"
                )

            response = await self.client.messages.create(
                model=self.model,
                max_tokens=self.max_tokens,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # Telemetry: LLM response
            llm_duration_ms = int((time.time() - llm_start_time) * 1000)
            input_tokens = response.usage.input_tokens
            output_tokens = response.usage.output_tokens
            cost_usd = (input_tokens / 1_000_000 * 3.0) + (output_tokens / 1_000_000 * 15.0)

            if TELEMETRY_AVAILABLE and self._current_trace_id and self._current_task_id:
                telemetry.llm_response(
                    provider="anthropic",
                    model=self.model,
                    trace_id=self._current_trace_id,
                    task_id=self._current_task_id,
                    duration_ms=llm_duration_ms,
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    cost_usd=cost_usd
                )

            content = response.content[0].text

            # Extract JSON (handles markdown code fences)
            import json
            json_str = self._extract_json_from_response(content)
            fields_data = json.loads(json_str)
            fields = fields_data.get('fields', [])

            logger.info(f"Extracted {len(fields)} fields")

            return fields

        except Exception as e:
            logger.error(f"Failed to extract fields: {e}")
            raise

    def _build_fields_prompt(
        self,
        doc_content: Dict[str, Any],
        structure: Dict[str, Any]
    ) -> str:
        """Build prompt for field extraction."""

        # Get sections that have fields
        field_sections = []
        for section in structure.get('sections', []):
            if section.get('has_fields'):
                field_sections.append(section)

        sections_text = "\n".join([f"- {s['title']}" for s in field_sections])

        # Format some paragraphs for context
        paragraphs_text = "\n\n".join([
            p['text'] for p in doc_content['paragraphs'][:50]
        ])

        prompt = f"""You are an expert at identifying fillable fields in ISO certification templates.

The document has these sections with fields:
{sections_text}

Document content (excerpt):
{paragraphs_text}

Your task:
Extract ALL fillable fields from the document, including:
- Form fields (text boxes, checkboxes, dropdowns)
- Questionnaire items
- Table cells that need to be filled
- Any placeholders like [Company Name], {{Date}}, etc.

For each field, determine:
1. Field name and label
2. Data type (text, number, date, email, phone, boolean, select, file)
3. Whether it's required or optional
4. Validation rules (max length, format, etc.)
5. Default value if any
6. Help text or description
7. Which section it belongs to

Return ONLY valid JSON in this exact format:
{{
  "fields": [
    {{
      "id": "company_name",
      "label": "Company Name",
      "section_id": "section_company_info",
      "type": "text",
      "required": true,
      "max_length": 200,
      "placeholder": "Enter legal company name",
      "help_text": "Official registered name of the organization",
      "validation_rules": {{
        "min_length": 2,
        "pattern": null
      }},
      "default_value": null
    }},
    {{
      "id": "certification_date",
      "label": "Target Certification Date",
      "section_id": "section_timeline",
      "type": "date",
      "required": true,
      "help_text": "Desired date for ISO certification audit",
      "validation_rules": {{
        "min_date": "today",
        "max_date": "+2years"
      }}
    }},
    {{
      "id": "has_quality_policy",
      "label": "Do you have a documented quality policy?",
      "section_id": "section_quality_system",
      "type": "boolean",
      "required": true,
      "help_text": "Required for ISO 9001 compliance"
    }}
  ]
}}

Important:
- Use snake_case for field IDs
- Choose the most appropriate type for each field
- Include helpful validation rules
- Match section_id to IDs from the structure
"""

        return prompt

    async def _validate_and_enrich(
        self,
        structure: Dict[str, Any],
        fields: List[Dict[str, Any]],
        iso_standard: Optional[str]
    ) -> Dict[str, Any]:
        """
        Validate extracted data and add enrichments.

        Returns:
            Complete template structure
        """
        # Build final template
        template = {
            "template_type": "iso_certification",
            "iso_standard": iso_standard or "ISO 9001:2015",
            "version": "1.0",
            "created_at": datetime.utcnow().isoformat(),
            "sections": structure.get('sections', []),
            "fields": fields,
            "metadata": {
                "total_sections": len(structure.get('sections', [])),
                "total_fields": len(fields),
                "required_fields": len([f for f in fields if f.get('required', False)]),
                "field_types": self._count_field_types(fields),
                "completion_estimate_minutes": self._estimate_completion_time(fields)
            },
            "validation": {
                "structure_valid": len(structure.get('sections', [])) > 0,
                "fields_valid": len(fields) > 0,
                "all_fields_have_sections": all(
                    any(f.get('section_id') == s['id'] for s in structure.get('sections', []))
                    for f in fields
                )
            }
        }

        logger.info(f"Template validated: {template['metadata']['total_sections']} sections, "
                   f"{template['metadata']['total_fields']} fields, "
                   f"~{template['metadata']['completion_estimate_minutes']} min to complete")

        return template

    def _count_field_types(self, fields: List[Dict[str, Any]]) -> Dict[str, int]:
        """Count fields by type."""
        counts = {}
        for field in fields:
            field_type = field.get('type', 'unknown')
            counts[field_type] = counts.get(field_type, 0) + 1
        return counts

    def _estimate_completion_time(self, fields: List[Dict[str, Any]]) -> int:
        """
        Estimate time to complete form (in minutes).

        Based on field complexity:
        - text: 1 min
        - number/date: 0.5 min
        - boolean: 0.25 min
        - select: 0.5 min
        - file: 2 min
        """
        time_by_type = {
            "text": 1.0,
            "number": 0.5,
            "date": 0.5,
            "email": 0.5,
            "phone": 0.5,
            "boolean": 0.25,
            "select": 0.5,
            "file": 2.0
        }

        total = 0
        for field in fields:
            field_type = field.get('type', 'text')
            total += time_by_type.get(field_type, 1.0)

        return int(total)
