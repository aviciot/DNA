"""
Template Agent - CORRECT APPROACH
===================================

AI-powered agent for parsing ISO policy documents.

New Approach:
- Identifies FIXED sections (policy text, same for all)
- Identifies FILLABLE sections (company-specific content)
- Semantic tagging of fillable sections
- Enables cross-document semantic mapping
"""

import json
import logging
import time
import os
from typing import Dict, Any, List, Optional
from pathlib import Path

from docx import Document
from agents.base_agent import BaseAgent
from template_validator import validate_template, ValidationError

logger = logging.getLogger(__name__)

# Configuration: Enable/disable self-healing
ENABLE_SELF_HEALING = os.getenv("ENABLE_TEMPLATE_SELF_HEALING", "true").lower() == "true"


class TemplateAgent(BaseAgent):
    """
    AI-powered agent for ISO template parsing.

    Key capability: Distinguish between fixed policy text
    and fillable company-specific sections.
    """

    @property
    def agent_name(self) -> str:
        """Agent name for logging/telemetry."""
        return "TemplateAgent"

    async def parse_document(
        self,
        file_path: str,
        custom_rules: Optional[str] = None,
        iso_standard: Optional[str] = None,
        trace_id: Optional[str] = None,
        task_id: Optional[str] = None,
        progress_callback: Optional[callable] = None
    ) -> Dict[str, Any]:
        """
        Parse ISO policy document - NEW APPROACH.

        Identifies:
        - Fixed sections (policy text)
        - Fillable sections (company-specific)

        Returns template structure with semantic tagging.
        """
        file_name = file_path.split('/')[-1]
        logger.info(f"Starting document parsing (NEW APPROACH): {file_name}")

        if trace_id:
            await self._start_operation(
                operation_name=f"Parse Template: {file_name}",
                trace_id=trace_id,
                task_id=task_id,
                file_path=file_path,
                iso_standard=iso_standard,
                has_custom_rules=bool(custom_rules)
            )

        start_time = time.time()

        try:
            # Step 1: Extract document content (40% of work)
            logger.info("Step 1/4: Loading Word document...")
            if progress_callback:
                await progress_callback(40, "Loading Word document...")
            doc_content = await self._extract_document_content(file_path)

            # Step 2: Identify fixed vs fillable sections (30% of work)
            logger.info("Step 2/4: Identifying fixed vs fillable sections...")
            if progress_callback:
                await progress_callback(70, "Analyzing document structure with AI...")
            template = await self._identify_sections(
                doc_content, iso_standard, custom_rules,
                trace_id=trace_id, task_id=task_id
            )

            # Step 3: Validate and self-heal if needed (15% of work)
            logger.info("Step 3/4: Validating template structure...")
            if progress_callback:
                await progress_callback(85, "Validating and self-healing template...")
            template = await self._validate_and_heal(
                template=template,
                doc_content=doc_content,
                iso_standard=iso_standard,
                custom_rules=custom_rules,
                trace_id=trace_id,
                task_id=task_id
            )

            # Step 4: Enrich with metadata (15% of work)
            logger.info("Step 4/4: Adding metadata...")
            if progress_callback:
                await progress_callback(95, "Finalizing template...")
            template = await self._enrich_template(template, file_name)

            logger.info(f"Parsing complete: {len(template['fixed_sections'])} fixed, "
                       f"{len(template['fillable_sections'])} fillable sections")

            # Telemetry: Complete
            duration = int(time.time() - start_time)
            if trace_id:
                await self._complete_operation(
                    operation_name=f"Parse Template: {file_name}",
                    trace_id=trace_id,
                    task_id=task_id,
                    duration_seconds=duration,
                    fixed_sections=len(template['fixed_sections']),
                    fillable_sections=len(template['fillable_sections'])
                )

            return template

        except Exception as e:
            if trace_id:
                await self._fail_operation(
                    operation_name=f"Parse Template: {file_name}",
                    trace_id=trace_id,
                    task_id=task_id,
                    error=str(e),
                    error_type="parsing_error"
                )
            raise

    # -------------------------------------------------------------------------
    # Document Extraction
    # -------------------------------------------------------------------------

    async def _extract_document_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from Word document.

        Returns paragraphs, tables, and metadata.
        """
        try:
            self._validate_document_file(file_path)

            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")

            doc = Document(file_path)

            # Extract paragraphs with structure
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal",
                        "level": self._get_heading_level(para)
                    })

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "Untitled",
                "author": core_props.author or "Unknown",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }

            logger.info(f"Extracted: {len(paragraphs)} paragraphs, {len(tables)} tables")

            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata
            }

        except FileNotFoundError:
            raise
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract document content: {e}")
            raise ValueError(f"Failed to read Word document: {str(e)}")

    def _get_heading_level(self, paragraph) -> int:
        """Get heading level from paragraph style."""
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if "heading 1" in style_name:
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        elif "heading" in style_name:
            try:
                return int(style_name.split()[-1])
            except:
                return 0
        return 0

    # -------------------------------------------------------------------------
    # NEW: Fixed vs Fillable Section Identification
    # -------------------------------------------------------------------------

    async def _identify_sections(
        self,
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str],
        trace_id: Optional[str] = None,
        task_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        NEW CORE METHOD: Identify fixed vs fillable sections.

        Returns:
        {
            "fixed_sections": [...],
            "fillable_sections": [...]
        }
        """
        # Build the prompt
        prompt = self._build_section_identification_prompt(doc_content, iso_standard, custom_rules)

        # Call LLM
        result = await self._call_llm(
            prompt=prompt,
            temperature=0.3,
            trace_id=trace_id,
            task_id=task_id,
            call_purpose="section_identification"
        )

        # Extract JSON
        json_str = self._extract_json(result["content"])

        # Parse JSON with error handling
        try:
            template = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed: {e}")
            logger.error(f"Failed JSON (first 500 chars): {json_str[:500]}")
            logger.error(f"Failed JSON (last 500 chars): {json_str[-500:]}")
            logger.error(f"JSON length: {len(json_str)} chars")

            # Try multiple repair strategies
            fixed_json = None
            repair_method = None

            # Strategy 1: Remove trailing commas
            try:
                fixed_json = json_str.replace(",]", "]").replace(",}", "}")
                template = json.loads(fixed_json)
                repair_method = "removed_trailing_commas"
            except json.JSONDecodeError:
                pass

            # Strategy 2: Try to close incomplete JSON (truncation)
            if repair_method is None:
                try:
                    fixed_json = json_str
                    # Count open/close brackets
                    open_braces = fixed_json.count("{") - fixed_json.count("}")
                    open_brackets = fixed_json.count("[") - fixed_json.count("]")

                    # Close missing brackets/braces
                    if open_braces > 0 or open_brackets > 0:
                        # Remove incomplete last entry (likely truncated)
                        # Find last complete field
                        if ',"' in fixed_json[-200:]:
                            last_comma = fixed_json.rfind(',"')
                            fixed_json = fixed_json[:last_comma]
                        elif '",{' in fixed_json[-200:]:
                            last_brace = fixed_json.rfind('",{')
                            fixed_json = fixed_json[:last_brace+2]

                        # Close arrays and objects
                        fixed_json += "]" * open_brackets
                        fixed_json += "}" * open_braces

                        template = json.loads(fixed_json)
                        repair_method = "closed_truncated_json"
                        logger.warning(f"JSON was truncated - closed {open_brackets} arrays and {open_braces} objects")
                except json.JSONDecodeError:
                    pass

            # Strategy 3: Extract valid portion up to error
            if repair_method is None:
                try:
                    # Try parsing up to the error position
                    valid_portion = json_str[:e.pos]
                    # Find last complete section
                    last_complete = valid_portion.rfind("}}")
                    if last_complete > 0:
                        valid_portion = valid_portion[:last_complete+2]
                        # Close arrays
                        valid_portion += "]}"
                        template = json.loads(valid_portion)
                        repair_method = "extracted_valid_portion"
                        logger.warning(f"Extracted valid JSON up to position {last_complete}")
                except json.JSONDecodeError:
                    pass

            if repair_method:
                logger.warning(f"✓ JSON repaired successfully using: {repair_method}")
            else:
                # All repair strategies failed
                raise ValueError(
                    f"LLM generated malformed JSON that cannot be repaired. Error: {e.msg} at position {e.pos}. "
                    f"Document has {doc_content['metadata']['paragraph_count']} paragraphs and is too complex. "
                    f"Try uploading a simpler document with fewer sections."
                )

        logger.info(f"Identified {len(template.get('fixed_sections', []))} fixed sections, "
                   f"{len(template.get('fillable_sections', []))} fillable sections")

        return template

    def _build_section_identification_prompt(
        self,
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str]
    ) -> str:
        """
        Build prompt for identifying fixed vs fillable sections.

        This is the CORE of the new approach!
        """
        # Format paragraphs
        paragraphs_text = "\n\n".join([
            f"[Level {p['level']}] {p['text']}"
            for p in doc_content['paragraphs'][:200]  # Include more context
        ])

        # Format tables
        tables_text = ""
        if doc_content['tables']:
            tables_text = f"\n\nTABLES FOUND: {len(doc_content['tables'])} tables\n"
            for i, table in enumerate(doc_content['tables'][:5]):
                tables_text += f"\nTable {i+1}:\n"
                for row in table[:3]:  # First 3 rows
                    tables_text += f"  {' | '.join(row)}\n"

        prompt = f"""You are an expert at analyzing ISO policy documents.

TASK: Identify FIXED vs FILLABLE sections in this document.

DOCUMENT METADATA:
- Title: {doc_content['metadata']['title']}
- Paragraphs: {doc_content['metadata']['paragraph_count']}
- Tables: {doc_content['metadata']['table_count']}
- ISO Standard: {iso_standard or 'Not specified'}

DOCUMENT CONTENT:
{paragraphs_text}
{tables_text}

{f'CUSTOM RULES: {custom_rules}' if custom_rules else ''}

---

INSTRUCTIONS:

Analyze this ISO policy document and categorize ALL content into:

1. **FIXED SECTIONS** (same for all companies):
   - Policy statements (General, Goal, Definition)
   - Standard procedures that don't change
   - Generic compliance requirements
   - Document control metadata tables (dates, versions, approvers)

   Example: "This policy applies to all parties operating within the company's network..."
   → This is FIXED - wording stays the same for everyone

2. **FILLABLE SECTIONS** (company-specific content):
   - "Relevant systems: _____"
   - Risk assessment tables with company details
   - "Our company uses _____"
   - Specific system names, technologies, processes
   - RTO/RPO values
   - Any content that varies per company

   Example: "Priority Level 1 - Relevant systems: Server DC & DevSRV"
   → This is FILLABLE - each company has different systems

For each FILLABLE section, identify:
- **location**: Where in document (section title, paragraph number)
- **type**: "table", "paragraph", "list", "field"
- **semantic_tags**: What kind of info is needed (e.g., ["infrastructure", "backup", "systems"])
- **current_content**: What's currently in the reference doc (as example)
- **format**: How content should be structured
- **is_mandatory**: true/false - Is this field required for compliance?
- **mandatory_confidence**: 0.0-1.0 - How confident are you this is mandatory?

DETECTING MANDATORY FIELDS:
Look for strong indicators near the field:
- HIGH CONFIDENCE (0.85-1.0): "mandatory", "required", "must be completed", "obligatory", "[MANDATORY]", "[REQUIRED]"
- MEDIUM CONFIDENCE (0.6-0.84): "must", "shall", "is required for"
- LOW CONFIDENCE (0.0-0.59): "should", "recommended" (DO NOT mark as mandatory)

ONLY mark is_mandatory=true if confidence >= 0.85 (be conservative - false positives are worse than false negatives)

Return ONLY valid JSON (no markdown):

{{
  "document_title": "ISMS XX Title",
  "fixed_sections": [
    {{
      "id": "general",
      "title": "General",
      "content": "This policy applies to all systems. [BE CONCISE - 1-2 sentences max]",
      "section_type": "policy_statement"
    }},
    {{
      "id": "goal",
      "title": "Goal",
      "content": "Ensure business continuity compliance. [SHORT]",
      "section_type": "policy_statement"
    }}
  ],
  "fillable_sections": [
    {{
      "id": "risk_assessment_table",
      "title": "Risk Assessment",
      "location": "Section 5.1",
      "type": "table",
      "semantic_tags": ["infrastructure", "hosting", "disaster-recovery", "cloud"],
      "current_content": "Example: Natural Disasters | Low | Our production servers...",
      "format": "Table with columns: Disaster type, Level of Impact, Anticipated Effect/Risk",
      "placeholder": "Describe your infrastructure, disaster recovery, and risk mitigation",
      "is_mandatory": true,
      "mandatory_confidence": 0.95
    }},
    {{
      "id": "priority_level_1_systems",
      "title": "Priority Level 1 - Relevant Systems",
      "location": "Section 5.6",
      "type": "list",
      "semantic_tags": ["critical-systems", "priorities", "infrastructure"],
      "current_content": "Server DC & DevSRV",
      "format": "Comma-separated list of system names",
      "placeholder": "List your most critical systems (RTO 0-8 hours)",
      "is_mandatory": false,
      "mandatory_confidence": 0.4
    }},
    {{
      "id": "backup_strategy",
      "title": "Data Backup And Recovery Options",
      "location": "Section 5.2",
      "type": "paragraph",
      "semantic_tags": ["backup", "disaster-recovery", "cloud", "storage"],
      "current_content": "Example text about backup procedures...",
      "format": "Paragraph describing backup approach",
      "placeholder": "Describe your backup and recovery strategy",
      "is_mandatory": true,
      "mandatory_confidence": 0.88
    }}
  ]
}}

IMPORTANT:
- Document control tables (dates, approvers) are FIXED structure (though content fills in)
- Be thorough - find ALL fillable spots
- Semantic tags help LLM map customer answers to sections
- One customer answer might fill multiple fillable sections!
- Current content shows what's in reference doc (as example)

JSON FORMATTING REQUIREMENTS:
- Return ONLY valid JSON (no markdown, no code fences, no explanations)
- BE CONCISE: Keep "content" fields to 2-3 sentences maximum
- FIXED sections: Just include first paragraph/sentence (not full text)
- FILLABLE sections: Describe what's needed, not full examples
- Ensure all arrays and objects are properly closed
- No trailing commas after last array/object elements
- All strings must be properly quoted
- Use double quotes (not single quotes)
- Validate JSON syntax before responding
- CRITICAL: If document has >20 sections, group similar policy sections together

RESPONSE SIZE LIMIT:
- Keep total response under 60KB (approximately 15K tokens)
- Prioritize fillable sections over fixed sections
- For fixed sections, just include title + first sentence
- Focus on COMPLETENESS over verbosity

Extract everything carefully. This is critical for semantic mapping!
"""

        return prompt

    # -------------------------------------------------------------------------
    # Template Validation & Self-Healing
    # -------------------------------------------------------------------------

    async def _validate_and_heal(
        self,
        template: Dict[str, Any],
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str],
        trace_id: Optional[str] = None,
        task_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Validate template structure and self-heal if needed.

        Steps:
        1. Run structural validation (required fields, types)
        2. If errors found and self-healing enabled, ask LLM to fix
        3. Re-validate healed template
        4. Run semantic validation (warnings only)

        Returns:
            Validated (and possibly healed) template

        Raises:
            ValueError: If validation fails after self-healing attempt
        """
        # Step 1: Structural validation
        errors, warnings = validate_template(template)

        # Log warnings (non-blocking)
        for warning in warnings:
            logger.warning(f"Template semantic issue: {warning}")

        # Step 2: If errors exist, try self-healing
        if errors:
            logger.warning(f"Template has {len(errors)} structural errors:")
            for error in errors:
                logger.warning(f"  - {error}")

            if ENABLE_SELF_HEALING:
                logger.info("Attempting self-healing with LLM...")
                try:
                    template = await self._self_heal_template(
                        original_template=template,
                        errors=errors,
                        doc_content=doc_content,
                        iso_standard=iso_standard,
                        custom_rules=custom_rules,
                        trace_id=trace_id,
                        task_id=task_id
                    )

                    # Step 3: Re-validate healed template
                    new_errors, new_warnings = validate_template(template)

                    if new_errors:
                        # Self-healing failed
                        logger.error(f"Self-healing failed, still have {len(new_errors)} errors:")
                        for error in new_errors:
                            logger.error(f"  - {error}")
                        raise ValueError(f"Template validation failed after self-heal: {[str(e) for e in new_errors]}")

                    logger.info("✅ Template self-healed successfully!")

                    # Log new warnings from healed template
                    for warning in new_warnings:
                        logger.warning(f"Template semantic issue (after heal): {warning}")

                except Exception as heal_error:
                    logger.error(f"Self-healing attempt failed: {heal_error}")
                    # Raise original validation errors
                    raise ValueError(f"Template validation failed: {[str(e) for e in errors]}")
            else:
                # Self-healing disabled
                logger.error("Self-healing is disabled, cannot fix template")
                raise ValueError(f"Template validation failed: {[str(e) for e in errors]}")

        return template

    async def _self_heal_template(
        self,
        original_template: dict,
        errors: List[ValidationError],
        doc_content: Dict[str, Any],
        iso_standard: Optional[str],
        custom_rules: Optional[str],
        trace_id: Optional[str] = None,
        task_id: Optional[str] = None
    ) -> dict:
        """
        Ask LLM to fix its own JSON output.

        This is a powerful pattern - the LLM is excellent at fixing its mistakes
        when you point them out clearly.

        Args:
            original_template: The template with errors
            errors: List of validation errors to fix
            doc_content: Original document content (for context)
            iso_standard: ISO standard (for context)
            custom_rules: Custom rules (for context)
            trace_id: Trace ID for telemetry
            task_id: Task ID for telemetry

        Returns:
            Healed template dictionary

        Raises:
            Exception: If LLM call fails or returns invalid JSON
        """
        # Format errors as a clear list
        error_list = "\n".join(f"  {i+1}. {error}" for i, error in enumerate(errors))

        # Build a focused healing prompt
        prompt = f"""You previously generated a template structure with some validation errors.

ORIGINAL OUTPUT (with errors):
```json
{json.dumps(original_template, indent=2)}
```

VALIDATION ERRORS FOUND:
{error_list}

INSTRUCTIONS:
Please fix ONLY the specific errors listed above.

Guidelines:
- Keep all existing section IDs unchanged
- Preserve all semantic tags
- Maintain all content that is correct
- Only fix the structural issues mentioned
- Ensure the output is valid JSON
- Do NOT add new sections or remove existing ones unless necessary to fix the errors

CONTEXT (for reference):
- Document: {doc_content['metadata']['title']}
- ISO Standard: {iso_standard or 'Not specified'}
{f"- Custom Rules: {custom_rules}" if custom_rules else ""}

Return ONLY the corrected JSON structure (no explanations, no markdown):"""

        logger.info("Calling LLM for self-healing...")

        # Call LLM with low temperature for precise fixes
        result = await self._call_llm(
            prompt=prompt,
            temperature=0.1,  # Low temperature for deterministic fixes
            trace_id=trace_id,
            task_id=task_id,
            call_purpose="self_heal_template"
        )

        # Extract and parse JSON
        json_str = self._extract_json(result["content"])

        # Use the same repair logic as initial parsing
        try:
            healed_template = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.error(f"Self-heal produced invalid JSON: {e}")
            # Try repair strategies
            try:
                fixed_json = json_str.replace(",]", "]").replace(",}", "}")
                healed_template = json.loads(fixed_json)
                logger.warning("Self-heal JSON repaired (removed trailing commas)")
            except json.JSONDecodeError:
                raise ValueError(f"Self-heal produced unparseable JSON at position {e.pos}")

        logger.info(f"Self-heal completed, template structure: "
                   f"{len(healed_template.get('fixed_sections', []))} fixed, "
                   f"{len(healed_template.get('fillable_sections', []))} fillable sections")

        return healed_template

    # -------------------------------------------------------------------------
    # Template Enrichment
    # -------------------------------------------------------------------------

    async def _enrich_template(
        self,
        template: Dict[str, Any],
        file_name: str
    ) -> Dict[str, Any]:
        """
        Enrich template with metadata and statistics.
        """
        # Calculate completion time estimate
        fillable_count = len(template.get("fillable_sections", []))
        completion_estimate_minutes = self._estimate_completion_time(fillable_count)

        # Add metadata
        template["metadata"] = {
            "source_file": file_name,
            "parsed_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "total_fixed_sections": len(template.get("fixed_sections", [])),
            "total_fillable_sections": fillable_count,
            "semantic_tags_used": self._extract_unique_tags(template.get("fillable_sections", [])),
            "completion_estimate_minutes": completion_estimate_minutes
        }

        return template

    def _estimate_completion_time(self, fillable_sections_count: int) -> int:
        """
        Estimate time to complete document (in minutes).

        Based on average time per fillable section:
        - Simple sections (field, paragraph): ~2 minutes
        - Complex sections (table, list): ~3 minutes

        Using conservative estimate of 2.5 minutes per section.
        """
        if fillable_sections_count == 0:
            return 5  # Minimum 5 minutes for review

        # 2.5 minutes per section, rounded up
        return max(5, int(fillable_sections_count * 2.5))

    def _extract_unique_tags(self, fillable_sections: List[Dict[str, Any]]) -> List[str]:
        """Extract unique semantic tags across all fillable sections."""
        tags = set()
        for section in fillable_sections:
            tags.update(section.get("semantic_tags", []))
        return sorted(list(tags))

    # -------------------------------------------------------------------------
    # Validation
    # -------------------------------------------------------------------------

    def _validate_document_file(self, file_path: str):
        """Validate document file before processing."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"Invalid file format: {path.suffix}. Only .docx/.doc supported.")

        # Check file size (50MB limit)
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > 50:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB. Maximum 50MB.")

        # Check readable
        try:
            with open(file_path, 'rb') as f:
                f.read(1024)
        except PermissionError:
            raise ValueError(f"Cannot read file: {file_path}")
        except Exception as e:
            raise ValueError(f"Cannot access file: {str(e)}")
