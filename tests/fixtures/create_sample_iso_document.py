"""
Create Sample ISO 9001 Document
================================

Generates a realistic ISO 9001 Quality Management System template
for testing the DNA parser.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_iso_9001_sample():
    """Create sample ISO 9001:2015 QMS document."""
    doc = Document()

    # Document properties
    doc.core_properties.title = "ISO 9001:2015 Quality Management System Application"
    doc.core_properties.author = "DNA Test Suite"

    # Title
    title = doc.add_heading('ISO 9001:2015 Quality Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Certification Application Form')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    doc.add_paragraph()

    # Section 1: Company Information
    doc.add_heading('1. Company Information', 1)
    doc.add_paragraph('Please provide your organization\'s basic information.')

    doc.add_paragraph('[Company Name]: _________________________________________')
    doc.add_paragraph('[Legal Entity Type]: Corporation / LLC / Partnership / Other')
    doc.add_paragraph('[Registration Number]: ___________________________________')
    doc.add_paragraph('[Founded Date]: __________________________________________')
    doc.add_paragraph('[Number of Employees]: ___________________________________')

    doc.add_heading('1.1 Contact Details', 2)
    doc.add_paragraph('[Primary Contact Name]: __________________________________')
    doc.add_paragraph('[Job Title]: _____________________________________________')
    doc.add_paragraph('[Email Address]: _________________________________________')
    doc.add_paragraph('[Phone Number]: __________________________________________')
    doc.add_paragraph('[Website]: _______________________________________________')

    doc.add_heading('1.2 Address', 2)
    doc.add_paragraph('[Street Address]: ________________________________________')
    doc.add_paragraph('[City]: __________________________________________________')
    doc.add_paragraph('[State/Province]: ________________________________________')
    doc.add_paragraph('[Postal Code]: ___________________________________________')
    doc.add_paragraph('[Country]: _______________________________________________')

    # Section 2: Business Scope
    doc.add_heading('2. Business Scope', 1)
    doc.add_paragraph('Describe your organization\'s scope of activities.')

    doc.add_paragraph('[Industry Sector]: ________________________________________')
    doc.add_paragraph('[Primary Products/Services]: ______________________________')
    doc.add_paragraph('[Secondary Products/Services]: ____________________________')

    doc.add_heading('2.1 Certification Scope', 2)
    doc.add_paragraph('Which processes do you want to include in the QMS certification?')
    doc.add_paragraph()

    # Checklist
    doc.add_paragraph('☐ Design and Development')
    doc.add_paragraph('☐ Production/Manufacturing')
    doc.add_paragraph('☐ Installation and Service')
    doc.add_paragraph('☐ Sales and Marketing')
    doc.add_paragraph('☐ Logistics and Distribution')

    # Section 3: Quality Management System
    doc.add_heading('3. Quality Management System Status', 1)

    doc.add_paragraph('Do you currently have a documented Quality Policy?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('Do you have documented Quality Objectives?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('Have you conducted internal audits in the last 12 months?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('[Date of Last Internal Audit]: ___________________________')
    doc.add_paragraph()

    doc.add_paragraph('Do you conduct management reviews?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('[Frequency of Management Reviews]: ________________________')

    # Section 4: Documentation
    doc.add_heading('4. Documentation', 1)
    doc.add_paragraph('Please indicate which documents you have in place:')
    doc.add_paragraph()

    doc.add_paragraph('☐ Quality Manual')
    doc.add_paragraph('☐ Procedures and Work Instructions')
    doc.add_paragraph('☐ Forms and Records')
    doc.add_paragraph('☐ Process Maps')
    doc.add_paragraph('☐ Organizational Chart')

    doc.add_heading('4.1 Document Upload', 2)
    doc.add_paragraph('[Quality Manual] (Upload PDF): __________________________')
    doc.add_paragraph('[Organization Chart] (Upload PDF): _______________________')
    doc.add_paragraph('[Process Maps] (Upload PDF): _____________________________')

    # Section 5: Timeline and Certification
    doc.add_heading('5. Certification Timeline', 1)

    doc.add_paragraph('[Target Certification Date]: _____________________________')
    doc.add_paragraph('[Preferred Audit Date]: __________________________________')
    doc.add_paragraph('[Number of Sites to be Certified]: _______________________')

    doc.add_heading('5.1 Previous Certifications', 2)
    doc.add_paragraph('Have you been certified to ISO 9001 before?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('[Previous Certification Body]: ____________________________')
    doc.add_paragraph('[Certificate Number]: _____________________________________')
    doc.add_paragraph('[Expiry Date]: ____________________________________________')

    # Section 6: Risk Management
    doc.add_heading('6. Risk Management', 1)

    doc.add_paragraph('Do you have a documented risk management process?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('Have you identified risks and opportunities?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('[Date of Last Risk Assessment]: ___________________________')

    # Section 7: Customer Satisfaction
    doc.add_heading('7. Customer Satisfaction', 1)

    doc.add_paragraph('Do you have a process for monitoring customer satisfaction?')
    doc.add_paragraph('☐ Yes  ☐ No')
    doc.add_paragraph()

    doc.add_paragraph('[Customer Satisfaction Score (%)]: ________________________')
    doc.add_paragraph('[Number of Customer Complaints (Last Year)]: ______________')

    # Section 8: Declaration
    doc.add_heading('8. Declaration', 1)
    doc.add_paragraph(
        'I declare that the information provided in this application is true and complete to the best of my knowledge.'
    )
    doc.add_paragraph()

    doc.add_paragraph('[Applicant Name]: _________________________________________')
    doc.add_paragraph('[Signature]: _____________________________________________')
    doc.add_paragraph('[Date]: __________________________________________________')

    # Save document
    output_path = 'tests/fixtures/sample_iso9001_qms_template.docx'
    doc.save(output_path)
    print(f"[OK] Sample ISO 9001 document created: {output_path}")
    print(f"  - 8 main sections")
    print(f"  - ~50 fillable fields")
    print(f"  - Multiple field types (text, date, boolean, select, file)")
    print(f"  - Realistic ISO 9001:2015 structure")

    return output_path


if __name__ == "__main__":
    create_iso_9001_sample()
